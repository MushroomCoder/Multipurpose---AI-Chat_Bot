# import os
# import json
# import pdfplumber
# from docx import Document
# import pandas as pd
# from llama_index.core import SimpleDirectoryReader, GPTListIndex

# def extract_pdf(file_path):
#     """Extract text from PDF."""
#     text = ""
#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             text += page.extract_text() + "\n"
#     return text.strip()

# def extract_docx(file_path):
#     """Extract text from DOCX."""
#     doc = Document(file_path)
#     text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
#     return text.strip()

# def extract_xlsx(file_path):
#     """Extract data from XLSX and convert to JSON."""
#     df = pd.read_excel(file_path)
#     return df.to_dict(orient="records")

# def extract_txt(file_path):
#     """Extract text from TXT."""
#     with open(file_path, "r", encoding="utf-8") as file:
#         return file.read().strip()

# def convert_to_json(file_path):
#     """Convert a document to a clean JSON format."""
#     ext = os.path.splitext(file_path)[1].lower()
#     if ext == ".pdf":
#         content = extract_pdf(file_path)
#     elif ext == ".docx":
#         content = extract_docx(file_path)
#     elif ext == ".xlsx":
#         content = extract_xlsx(file_path)
#     elif ext == ".txt":
#         content = extract_txt(file_path)
#     else:
#         raise ValueError(f"Unsupported file format: {ext}")

#     return {"file_name": os.path.basename(file_path), "content": content}

# def create_index_from_json(json_data):
#     """Create an index using LlamaIndex."""
#     # Save JSON data temporarily
#     temp_dir = "temp_docs"
#     os.makedirs(temp_dir, exist_ok=True)
#     for idx, doc in enumerate(json_data):
#         with open(os.path.join(temp_dir, f"doc_{idx}.txt"), "w") as f:
#             f.write(json.dumps(doc))

#     # Use SimpleDirectoryReader to read documents
#     documents = SimpleDirectoryReader(temp_dir).load_data()
#     index = GPTListIndex.from_documents(documents)

#     # Clean up temporary files
#     for file in os.listdir(temp_dir):
#         os.remove(os.path.join(temp_dir, file))
#     os.rmdir(temp_dir)

#     return index

# # Main function to handle a single document
# def process_document(file_path):
#     try:
#         # Convert document to JSON
#         json_output = [convert_to_json(file_path)]

#         # Save JSON to a file
#         json_file = os.path.splitext(file_path)[0] + ".json"
#         with open(json_file, "w", encoding="utf-8") as f:
#             json.dump(json_output, f, indent=4)

#         print(f"JSON output saved to {json_file}")

#         # Create LlamaIndex
#         index = create_index_from_json(json_output)
#         print("Index created successfully!")

#         return index
#     except Exception as e:
#         print(f"Error processing document: {e}")

# # Example usage
# if __name__ == "__main__":
#     file_path = input("Enter the path of the document: ").strip()
#     if os.path.exists(file_path):
#         process_document(file_path)
#     else:
#         print("File does not exist. Please check the path and try again.")

import os
import json
from llama_index.core import SimpleDirectoryReader, GPTVectorStoreIndex
from docx import Document
import pandas as pd
from PyPDF2 import PdfReader

def read_file(file_path):
    """Read content from a given file based on its type."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return read_pdf(file_path)
    elif ext == '.docx':
        return read_docx(file_path)
    elif ext == '.txt':
        return read_txt(file_path)
    elif ext == '.xlsx':
        return read_xlsx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def read_pdf(file_path):
    """Extract text from a PDF file."""
    reader = PdfReader(file_path)
    content = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    return content

def read_docx(file_path):
    """Extract text from a DOCX file."""
    doc = Document(file_path)
    content = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
    return content

def read_txt(file_path):
    """Read text content from a TXT file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    return content

def read_xlsx(file_path):
    """Extract text from an XLSX file."""
    data = pd.read_excel(file_path, sheet_name=None)
    content = ""
    for sheet, df in data.items():
        content += f"Sheet: {sheet}\n"
        content += df.to_string(index=False, header=True)
        content += "\n"
    return content

def convert_to_json(file_path, output_path):
    """Convert the document content into a structured JSON format."""
    try:
        content = read_file(file_path)
        structured_data = {
            "file_name": os.path.basename(file_path),
            "content": content.strip()
        }
        with open(output_path, 'w', encoding='utf-8') as json_file:
            json.dump(structured_data, json_file, indent=4, ensure_ascii=False)
        print(f"File successfully converted to JSON: {output_path}")
        return structured_data
    except Exception as e:
        print(f"Error: {e}")
        return None

def build_index_from_json(json_data):
    """Build an index using LlamaIndex from JSON data."""
    with open("temp_text.txt", 'w', encoding='utf-8') as temp_file:
        temp_file.write(json_data["content"])
    documents = SimpleDirectoryReader(input_dir=".", input_files=["temp_text.txt"]).load_data()
    index = GPTVectorStoreIndex.from_documents(documents)
    os.remove("temp_text.txt")
    return index

def main():
    file_path = input("Enter the file path: ").strip()
    output_path = "output.json"

    # Convert document to JSON
    json_data = convert_to_json(file_path, output_path)

    if json_data:
        # Build index from JSON
        index = build_index_from_json(json_data)
        print("Index built successfully!")

if __name__ == "__main__":
    main()
